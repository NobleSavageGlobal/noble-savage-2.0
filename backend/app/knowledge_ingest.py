from __future__ import annotations

from dataclasses import dataclass
import hashlib
from io import BytesIO
from pathlib import Path
import re
from typing import Any


MAX_CHUNK_CHARS = 9000
MAX_UPLOAD_BYTES = 20 * 1024 * 1024

TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".csv",
    ".json",
    ".yaml",
    ".yml",
    ".xml",
}

CODE_EXTENSIONS = {
    ".py",
    ".js",
    ".ts",
    ".tsx",
    ".jsx",
    ".java",
    ".go",
    ".rs",
    ".rb",
    ".php",
    ".cs",
    ".cpp",
    ".c",
    ".h",
    ".hpp",
    ".swift",
    ".kt",
    ".sql",
    ".sh",
    ".html",
    ".css",
    ".scss",
}

IMAGE_EXTENSIONS = {
    ".png",
    ".jpg",
    ".jpeg",
    ".webp",
    ".bmp",
    ".tif",
    ".tiff",
}


@dataclass
class ParsedDocument:
    title: str
    source: str
    content: str
    tags: list[str]
    warnings: list[str]
    token_count: int = 0
    chunk_count: int = 0
    file_id: str = ""
    ocr_used: bool = False


def _normalize_whitespace(value: str) -> str:
    compact = value.replace("\x00", "")
    compact = re.sub(r"\r\n?", "\n", compact)
    compact = re.sub(r"\n{3,}", "\n\n", compact)
    return compact.strip()


def _split_sentences(text: str) -> list[str]:
    """Split text into sentence-like segments, preferring sentence boundaries over raw positions."""
    return re.split(r"(?<=[.!?])\s+(?=[A-Z\u201c\u2018\"])", text)


def _split_chunks(text: str, max_chars: int = MAX_CHUNK_CHARS) -> list[str]:
    cleaned = _normalize_whitespace(text)
    if not cleaned:
        return []
    if len(cleaned) <= max_chars:
        return [cleaned]

    # First split on paragraph breaks
    paragraphs = re.split(r"\n\n+", cleaned)
    chunks: list[str] = []
    current = ""

    for para in paragraphs:
        candidate = f"{current}\n\n{para}".strip() if current else para
        if len(candidate) <= max_chars:
            current = candidate
            continue

        # Paragraph fits in one chunk but combined doesn't — flush current
        if current:
            chunks.append(current)
            current = ""

        # Paragraph itself is too large — split on sentence boundaries
        if len(para) > max_chars:
            sentences = _split_sentences(para)
            for sentence in sentences:
                sentence = sentence.strip()
                if not sentence:
                    continue
                candidate = f"{current} {sentence}".strip() if current else sentence
                if len(candidate) <= max_chars:
                    current = candidate
                else:
                    if current:
                        chunks.append(current)
                    # Single sentence longer than max — hard split as last resort
                    while len(sentence) > max_chars:
                        chunks.append(sentence[:max_chars])
                        sentence = sentence[max_chars:]
                    current = sentence
        else:
            current = para

    if current:
        chunks.append(current)
    return chunks


def _ocr_image_bytes(image_bytes: bytes) -> str:
    try:
        from PIL import Image
        import pytesseract
    except Exception:
        return ""

    try:
        image = Image.open(BytesIO(image_bytes))
        return pytesseract.image_to_string(image) or ""
    except Exception:
        return ""


def _extract_pdf_with_ocr(reader: Any) -> str:
    extracted: list[str] = []
    for page in reader.pages:
        for image in getattr(page, "images", []):
            text = _ocr_image_bytes(image.data)
            if text.strip():
                extracted.append(text)
    return "\n\n".join(extracted)


def parse_document(name: str, content_type: str | None, data: bytes) -> ParsedDocument:
    if not data:
        raise ValueError("File is empty.")
    if len(data) > MAX_UPLOAD_BYTES:
        raise ValueError("File exceeds 20MB upload limit.")

    suffix = Path(name).suffix.lower()
    stem = Path(name).stem.strip() or "Imported document"
    source = name
    warnings: list[str] = []
    ocr_used = False

    if suffix == ".pdf":
        try:
            from pypdf import PdfReader
        except Exception as exc:
            raise ValueError("PDF support is not available on this server.") from exc

        reader = PdfReader(BytesIO(data))
        pages = []
        for page in reader.pages:
            pages.append(page.extract_text() or "")
        content = "\n\n".join(pages)
        if len(_normalize_whitespace(content)) < 80:
            ocr_text = _extract_pdf_with_ocr(reader)
            if ocr_text.strip():
                content = ocr_text
                ocr_used = True
                warnings.append("Text layer was sparse; OCR fallback was used for scanned content.")
            else:
                warnings.append("PDF appears image-based and OCR extraction could not recover readable text.")
        tags = ["pdf", "imported"]
    elif suffix == ".docx":
        try:
            from docx import Document
        except Exception as exc:
            raise ValueError("DOCX support is not available on this server.") from exc

        document = Document(BytesIO(data))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        content = "\n".join(paragraphs)
        tags = ["docx", "imported"]
    elif suffix in IMAGE_EXTENSIONS:
        content = _ocr_image_bytes(data)
        ocr_used = bool(content.strip())
        if not ocr_used:
            warnings.append("Image OCR returned no readable text.")
        tags = ["image", "ocr", "imported"]
    elif suffix in TEXT_EXTENSIONS or suffix in CODE_EXTENSIONS:
        content = data.decode("utf-8", errors="ignore")
        tags = ["code", "imported"] if suffix in CODE_EXTENSIONS else ["text", "imported"]
    elif content_type and content_type.startswith("text/"):
        content = data.decode("utf-8", errors="ignore")
        tags = ["text", "imported"]
    else:
        raise ValueError(f"Unsupported file type: {suffix or content_type or 'unknown'}")

    normalized = _normalize_whitespace(content)
    if not normalized:
        raise ValueError("Could not extract readable text from file.")
    token_count = len(re.findall(r"\S+", normalized))
    chunk_count = len(_split_chunks(normalized))
    file_id = hashlib.sha1(f"{name}:{len(data)}:{normalized[:300]}".encode("utf-8")).hexdigest()

    return ParsedDocument(
        title=stem[:180],
        source=source[:280],
        content=normalized,
        tags=tags,
        warnings=warnings,
        token_count=token_count,
        chunk_count=chunk_count,
        file_id=file_id,
        ocr_used=ocr_used,
    )


def build_knowledge_payloads(parsed: ParsedDocument) -> list[dict[str, str | list[str]]]:
    chunks = _split_chunks(parsed.content)
    payloads: list[dict[str, str | list[str]]] = []
    total = len(chunks)
    for index, chunk in enumerate(chunks, start=1):
        title = parsed.title if total == 1 else f"{parsed.title} (Part {index}/{total})"
        payloads.append(
            {
                "title": title[:200],
                "content": chunk,
                "source": parsed.source,
                "tags": parsed.tags,
                "file_id": parsed.file_id,
                "chunk_index": index,
                "chunk_total": total,
                "token_count": len(re.findall(r"\S+", chunk)),
                "status": "indexed",
            }
        )
    return payloads
